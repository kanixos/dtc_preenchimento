import streamlit as st
import pandas as pd
from docx import Document
import io
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
import tempfile
import os

# ==========================================
# FUNÇÕES AUXILIARES DE FORMATAÇÃO
# ==========================================
def encontrar_paragrafo(doc, texto_tag):
    for p in doc.paragraphs:
        if texto_tag in p.text: return p
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if texto_tag in p.text: return p
    return None

def preencher_celula(celula, texto, fundo_cinza=False, negrito=False, alinhar_centro=True, espacamento=False):
    # Se espacamento for True, injeta as quebras de linha (3 espaços) para a tabela de remuneração
    if espacamento:
        celula.text = f"\n{texto}\n"
    else:
        celula.text = str(texto)
    
    # Centralização Vertical
    celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Formatação do Parágrafo e Fonte
    for p in celula.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alinhar_centro else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = negrito
            
    if fundo_cinza:
        tcPr = celula._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)

def definir_larguras(tabela, larguras):
    """Ajusta a largura rígida de cada coluna em Centímetros para garantir alinhamento perfeito."""
    tabela.autofit = False
    for row in tabela.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(larguras):
                cell.width = larguras[idx]

def substituir_tags_preservando_formatacao(doc, dados_pessoais):
    """Substitui as tags de forma segura mantendo negritos e tamanhos específicos."""
    dtc_count = 0
    for p in doc.paragraphs:
        for chave, valor in dados_pessoais.items():
            if chave in p.text:
                if chave == "{{NUM_DTC}}":
                    dtc_count += 1
                    p.text = p.text.replace(chave, str(valor))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.bold = True
                        if dtc_count == 1:
                            run.font.size = Pt(14) # Primeira página (Maior)
                        else:
                            run.font.size = Pt(12) # Segunda página (Remunerações)
                else:
                    p.text = p.text.replace(chave, str(valor))
                    # Fallback para outras tags soltas fora de tabelas
                    for run in p.runs:
                        run.font.name = 'Calibri'

    # Substituição nas Tabelas (Dados Pessoais)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for chave, valor in dados_pessoais.items():
                        if chave in p.text:
                            # Tenta substituir direto no run para preservar negrito das labels (ex: "NOME:")
                            substituido = False
                            for run in p.runs:
                                if chave in run.text:
                                    run.text = run.text.replace(chave, str(valor))
                                    run.font.name = 'Calibri'
                                    substituido = True
                            if not substituido:
                                p.text = p.text.replace(chave, str(valor))
                                for run in p.runs:
                                    run.font.name = 'Calibri'

# ==========================================
# MOTOR INTELIGENTE DE PREENCHIMENTO
# ==========================================
def preencher_documento(dados_pessoais, df_vinculos, df_remuneracoes):
    doc = Document("template.docx")
    
    substituir_tags_preservando_formatacao(doc, dados_pessoais)

    # =========================================================
    # TABELA 1: DADOS FUNCIONAIS (1ª Página)
    # =========================================================
    p_func1 = encontrar_paragrafo(doc, '{{TAB_FUNC_1}}')
    if p_func1 is not None:
        tbl1 = doc.add_table(rows=0, cols=3)
        tbl1.style = 'Table Grid'
        
        for idx, row in df_vinculos.iterrows():
            if pd.isna(row.iloc[0]) or str(row.iloc[0]).strip() == "": continue
            v = str(row['Vínculo'])
            
            r_adm = tbl1.add_row().cells
            preencher_celula(r_adm[0], f"DATA DE ADMISSÃO NO VÍNCULO {v}:\n{row.get('Dt. Admissão', '-')}", alinhar_centro=False)
            preencher_celula(r_adm[1], f"Nº DE PORTARIA DE NOMEAÇÃO:\n{row.get('Port. Nomeação', 'NA')}", alinhar_centro=False)
            preencher_celula(r_adm[2], f"DATA DA PUBLICAÇÃO:\n{row.get('Pub. Nomeação', 'NA')}", alinhar_centro=False)
            
            r_des = tbl1.add_row().cells
            preencher_celula(r_des[0], f"DATA DE DESLIGAMENTO NO VÍNCULO {v}:\n{row.get('Dt. Desligamento', '-')}", alinhar_centro=False)
            preencher_celula(r_des[1], f"Nº DE PORTARIA DE EXONERAÇÃO/ DEMISSÃO:\n{row.get('Port. Exoneração', 'NA')}", alinhar_centro=False)
            preencher_celula(r_des[2], f"DATA DA PUBLICAÇÃO:\n{row.get('Pub. Exoneração', 'NA')}", alinhar_centro=False)
        
        # Total Exato: 16.0 cm (Alinhado com as margens de 2.5cm)
        definir_larguras(tbl1, [Cm(5.33), Cm(5.33), Cm(5.34)])
        p_func1._p.addnext(tbl1._tbl)
        p_func1.text = p_func1.text.replace('{{TAB_FUNC_1}}', '')

    # =========================================================
    # TABELA 2: PERÍODOS DE CONTRIBUIÇÃO (1ª Página)
    # =========================================================
    p_per = encontrar_paragrafo(doc, '{{TAB_PER}}')
    if p_per is not None:
        tbl2 = doc.add_table(rows=1, cols=5)
        tbl2.style = 'Table Grid'
        
        headers = ["SEQ.:", "DATA INÍCIO:", "DATA FIM:", "CARGO/FUNÇÃO:", "CATEGORIA FUNCIONAL:"]
        for i, h in enumerate(headers):
            preencher_celula(tbl2.cell(0, i), h, negrito=True)
        
        for idx, row in df_vinculos.iterrows():
            if pd.isna(row.iloc[0]) or str(row.iloc[0]).strip() == "": continue
            seq = str(idx + 1).zfill(2)
            r = tbl2.add_row().cells
            
            cat = str(row.get('Categoria Funcional', ''))
            cat_efet = "X" if cat == "Efetivo/Estável" else " "
            cat_comis = "X" if cat == "Comissionado/Mandato Eletivo" else " "
            cat_contrat = "X" if cat == "Contratado" else " "
            cat_texto = f"({cat_efet}) Efetivo/Estável\n({cat_comis}) Comissionado/Mandato Eletivo\n({cat_contrat}) Contratado"

            preencher_celula(r[0], seq)
            preencher_celula(r[1], str(row.get('Dt. Admissão', '-')))
            preencher_celula(r[2], str(row.get('Dt. Desligamento', '-')))
            preencher_celula(r[3], str(row.get('Cargo / Função', '-')).upper())
            preencher_celula(r[4], cat_texto, alinhar_centro=False)

        if len(tbl2.rows) > 1:
            for col_idx in [3, 4]:
                start_row = 1
                while start_row < len(tbl2.rows):
                    end_row = start_row
                    while end_row + 1 < len(tbl2.rows) and tbl2.cell(end_row + 1, col_idx).text == tbl2.cell(start_row, col_idx).text:
                        end_row += 1
                    if end_row > start_row:
                        start_cell = tbl2.cell(start_row, col_idx)
                        raw_text = start_cell.text.strip()
                        for r_idx in range(start_row + 1, end_row + 1):
                            start_cell.merge(tbl2.cell(r_idx, col_idx))
                        preencher_celula(start_cell, raw_text, alinhar_centro=(col_idx==3)) 
                    start_row = end_row + 1

        # Total Exato: 16.0 cm
        definir_larguras(tbl2, [Cm(1.5), Cm(3.0), Cm(3.0), Cm(4.0), Cm(4.5)])
        p_per._p.addnext(tbl2._tbl)
        p_per.text = p_per.text.replace('{{TAB_PER}}', '')

    # =========================================================
    # TABELA 3: DADOS FUNCIONAIS 2 (2ª Página)
    # =========================================================
    p_func2 = encontrar_paragrafo(doc, '{{TAB_FUNC_2}}')
    if p_func2 is not None:
        tbl3 = doc.add_table(rows=0, cols=4)
        tbl3.style = 'Table Grid'
        
        for idx, row in df_vinculos.iterrows():
            if pd.isna(row.iloc[0]) or str(row.iloc[0]).strip() == "": continue
            v = str(row['Vínculo'])
            r = tbl3.add_row().cells
            
            preencher_celula(r[0], f"DATA DE ADMISSÃO\nNO VÍNCULO {v}:\n{row.get('Dt. Admissão', '-')}", alinhar_centro=False)
            preencher_celula(r[1], f"DATA DE EXONERAÇÃO\nNO VÍNCULO {v}:\n{row.get('Dt. Desligamento', '-')}", alinhar_centro=False)
            
            if idx == 0:
                preencher_celula(r[2], f"PIS/PASEP:\n{dados_pessoais.get('{{PIS}}', '')}", alinhar_centro=False)
                preencher_celula(r[3], f"CPF:\n{dados_pessoais.get('{{CPF}}', '')}", alinhar_centro=False)
            else:
                preencher_celula(r[2], "", alinhar_centro=False)
                preencher_celula(r[3], "", alinhar_centro=False)
        
        if len(tbl3.rows) > 1:
            for col_idx in [2, 3]:
                start_cell = tbl3.cell(0, col_idx)
                raw_text = start_cell.text.strip()
                for r_idx in range(1, len(tbl3.rows)):
                    start_cell.merge(tbl3.cell(r_idx, col_idx))
                preencher_celula(start_cell, raw_text, alinhar_centro=False)

        # Total Exato: 16.0 cm
        definir_larguras(tbl3, [Cm(4.0), Cm(4.0), Cm(4.0), Cm(4.0)])
        p_func2._p.addnext(tbl3._tbl)
        p_func2.text = p_func2.text.replace('{{TAB_FUNC_2}}', '')

    # =========================================================
    # TABELA 4: MATRIZ DE REMUNERAÇÕES PADRÃO INSS (Com Espaçamento)
    # =========================================================
    p_remun = encontrar_paragrafo(doc, '{{TABELAS_REMUNERACAO}}')
    if p_remun is not None:
        anos_encontrados = set()
        dados_matriz = {}
        
        anos_colunas = [c for c in df_remuneracoes.columns if c != 'Mês']
        
        for row_idx, row in df_remuneracoes.iterrows():
            mes_num = row_idx + 1 # 1 para Janeiro, 2 para Fevereiro...
            for ano_str in anos_colunas:
                try:
                    ano_int = int(ano_str)
                    anos_encontrados.add(ano_int)
                    val = str(row.get(ano_str, '')).strip()
                    if val and val.lower() not in ['nan', 'none', '<na>']:
                        dados_matriz[(mes_num, ano_int)] = val
                except ValueError:
                    pass

        # GERA UMA LINHA DE ANOS CONTÍNUA E FORÇA A TER MÚLTIPLOS DE 5 ANOS
        if anos_encontrados:
            min_ano = min(anos_encontrados)
            max_ano = max(anos_encontrados)
            anos_completos = list(range(min_ano, max_ano + 1))
            
            # Preenche com anos extras vazios para fechar blocos de exatos 5 anos
            while len(anos_completos) % 5 != 0:
                anos_completos.append(anos_completos[-1] + 1)
        else:
            anos_completos = []

        blocos_de_anos = [anos_completos[i:i+5] for i in range(0, len(anos_completos), 5)]
        meses_nomes = ["JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO", "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"]

        elemento_anterior = p_remun._p
        
        for bloco in blocos_de_anos:
            tbl = doc.add_table(rows=14, cols=len(bloco)+1)
            tbl.style = 'Table Grid'
            
            c_mes0 = tbl.cell(0, 0)
            c_mes0.merge(tbl.cell(1, 0))
            # Ativando 'espacamento=True' para dar os 3 espaços de altura
            preencher_celula(c_mes0, "Mês", fundo_cinza=True, negrito=True, espacamento=True)
            
            for col_idx, ano in enumerate(bloco):
                c_ano = tbl.cell(0, col_idx+1)
                preencher_celula(c_ano, f"Ano: {ano}", fundo_cinza=True, negrito=True, espacamento=True)
                
                c_val = tbl.cell(1, col_idx+1)
                # "Valor($)" sem negrito conforme pedido!
                preencher_celula(c_val, "Valor($)", fundo_cinza=True, negrito=False, espacamento=True)
                
            for mes_idx, mes_nome in enumerate(meses_nomes):
                c_m = tbl.cell(mes_idx+2, 0)
                # Meses em negrito
                preencher_celula(c_m, mes_nome, negrito=True, espacamento=True)
                
                for col_idx, ano in enumerate(bloco):
                    val = dados_matriz.get((mes_idx+1, ano), "-")
                    c_v = tbl.cell(mes_idx+2, col_idx+1)
                    preencher_celula(c_v, val, espacamento=True)
            
            # Total Exato: 16.0 cm (Mês=3.5cm, Cada Ano=2.5cm * 5)
            larguras_t4 = [Cm(3.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)]
            definir_larguras(tbl, larguras_t4)
            
            elemento_anterior.addnext(tbl._tbl)
            elemento_anterior = tbl._tbl
            
            p_space = OxmlElement('w:p')
            elemento_anterior.addnext(p_space)
            elemento_anterior = p_space
            
        p_remun.text = p_remun.text.replace('{{TABELAS_REMUNERACAO}}', '')

    arquivo_gerado = io.BytesIO()
    doc.save(arquivo_gerado)
    arquivo_gerado.seek(0)
    return arquivo_gerado

# ==========================================
# INTERFACE DO STREAMLIT
# ==========================================
st.set_page_config(page_title="Preenchedor Automático DTC", layout="wide")

st.title("📄 Preenchedor Automático DTC")
st.write("Interface otimizada. O sistema desenhará magicamente as tabelas exatas do padrão INSS com fontes e cores corretas.")

# --- 1. Dados Pessoais ---
st.subheader("1. Identificação do Documento e Servidor")

col_dtc1, col_dtc2 = st.columns([1, 3])
with col_dtc1:
    numero_dtc = st.text_input("Número do DTC", placeholder="Ex: 199/2025")

st.markdown("---")

col1, col2, col3 = st.columns(3)
with col1:
    nome = st.text_input("Nome do Servidor")
    rg = st.text_input("RG / Órgão Emissor")
    mae = st.text_input("Nome da Mãe")
    pai = st.text_input("Nome do Pai")
with col2:
    cpf = st.text_input("CPF")
    pis = st.text_input("PIS/PASEP")
with col3:
    matricula = st.text_input("Matrícula")
    data_nasc = st.text_input("Data de Nascimento")

dados_pessoais = {
    "{{NUM_DTC}}": numero_dtc,
    "{{NOME}}": nome, "{{CPF}}": cpf, "{{MATRICULA}}": matricula,
    "{{RG}}": rg, "{{PIS}}": pis, "{{MAE}}": mae, "{{PAI}}": pai,
    "{{DATA_NASC}}": data_nasc,
    "{{DATA}}": datetime.today().strftime('%d/%m/%Y')
}

st.markdown("---")

# --- 2. Vínculos Funcionais ---
st.subheader("2. Vínculos, Períodos e Cargos")

colunas_vinc = [
    "Vínculo", "Dt. Admissão", "Dt. Desligamento", 
    "Cargo / Função", "Categoria Funcional", 
    "Port. Nomeação", "Pub. Nomeação", "Port. Exoneração", "Pub. Exoneração"
]
df_vinc_vazio = pd.DataFrame(columns=colunas_vinc, index=range(3))

config_colunas = {
    "Categoria Funcional": st.column_config.SelectboxColumn(
        "Categoria Funcional",
        options=["Efetivo/Estável", "Comissionado/Mandato Eletivo", "Contratado"],
        required=True
    )
}

df_vinculos = st.data_editor(df_vinc_vazio, num_rows="dynamic", use_container_width=True, column_config=config_colunas)

st.markdown("---")

# --- 3. Remunerações ---
st.subheader("3. Remunerações")
st.success("⚡ **Acelerador Visual:** Digite os salários diretamente na matriz abaixo, do mesmo jeito que eles vão aparecer no documento final!")

col_ano1, col_ano2, col_btn = st.columns([1, 1, 2])
with col_ano1:
    ano_inicio = st.number_input("Ano Inicial", min_value=1990, max_value=2050, value=2018, step=1)
with col_ano2:
    ano_fim = st.number_input("Ano Final", min_value=1990, max_value=2050, value=2022, step=1)

meses_nomes_ui = ["JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO", "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"]

# Inicializa a matriz na memória se não existir
if 'df_remun_matriz' not in st.session_state:
    colunas_padrao = ['Mês'] + [str(ano) for ano in range(2018, 2023)]
    df_padrao = pd.DataFrame(columns=colunas_padrao)
    df_padrao['Mês'] = meses_nomes_ui
    st.session_state.df_remun_matriz = df_padrao

with col_btn:
    st.write("") 
    if st.button("⬇️ Gerar Tabela de Anos", type="secondary", use_container_width=True):
        colunas_novas = ['Mês'] + [str(ano) for ano in range(ano_inicio, ano_fim + 1)]
        df_nova = pd.DataFrame(columns=colunas_novas)
        df_nova['Mês'] = meses_nomes_ui
        st.session_state.df_remun_matriz = df_nova

# Configura a coluna Mês para ser um "índice fixo" (não editável) no visual do Streamlit
config_colunas_remun = {"Mês": st.column_config.TextColumn("Mês", disabled=True)}

df_remuneracoes = st.data_editor(
    st.session_state.df_remun_matriz, 
    hide_index=True, 
    use_container_width=True, 
    column_config=config_colunas_remun,
    height=450
)

# --- Botão de Geração Final ---
st.markdown("---")

gerar_pdf = st.checkbox("Gerar também em formato PDF (Requer execução local no Windows com MS Word instalado)")

if st.button("🚀 GERAR DOCUMENTO PADRÃO INSS", type="primary", use_container_width=True):
    if nome and cpf:
        with st.spinner('A desenhar as tabelas exatas do INSS com fontes e cores...'):
            try:
                arquivo_docx = preencher_documento(dados_pessoais, df_vinculos, df_remuneracoes)
                st.balloons()
                
                col_down1, col_down2 = st.columns(2)
                
                with col_down1:
                    st.download_button(
                        label="📥 BAIXAR DTC (WORD)",
                        data=arquivo_docx,
                        file_name=f"DTC_{nome.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                if gerar_pdf:
                    with col_down2:
                        try:
                            from docx2pdf import convert
                            
                            with st.spinner("Convertendo para PDF (isso pode levar alguns segundos)..."):
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                                    tmp_docx.write(arquivo_docx.getvalue())
                                    tmp_docx_path = tmp_docx.name
                                    
                                tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")
                                
                                convert(tmp_docx_path, tmp_pdf_path)
                                
                                with open(tmp_pdf_path, "rb") as f:
                                    pdf_bytes = f.read()
                                    
                                st.download_button(
                                    label="📥 BAIXAR DTC (PDF)",
                                    data=pdf_bytes,
                                    file_name=f"DTC_{nome.replace(' ', '_')}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                                
                                try:
                                    os.remove(tmp_docx_path)
                                    os.remove(tmp_pdf_path)
                                except:
                                    pass
                                    
                        except ImportError:
                            st.error("⚠️ Biblioteca 'docx2pdf' ausente. Abra o terminal e digite: pip install docx2pdf")
                        except Exception as e:
                            st.warning("⚠️ O sistema de PDF não funcionou na nuvem. Baixe o Word e vá em 'Salvar como PDF' no seu computador.")
                            
            except Exception as e:
                st.error(f"Erro ao processar as matrizes. Detalhe técnico: {e}")
    else:
        st.error("⚠️ Preencha pelo menos o Nome e o CPF do servidor.")

# --- Rodapé ---
st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: gray; font-size: 12px;'><i>Desenvolvido por André Almeida Costa</i></p>", unsafe_allow_html=True)
